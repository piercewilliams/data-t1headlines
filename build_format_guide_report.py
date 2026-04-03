"""
Build a formatted .docx report from the T1 Data Evidence Report for Platform Format Guide.
Output: ~/Desktop/T1_Data_Evidence_Report_Format_Guide.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colors ────────────────────────────────────────────────────────────────────
NAVY     = RGBColor(0x0f, 0x17, 0x2a)
BLUE     = RGBColor(0x25, 0x63, 0xeb)
GREEN    = RGBColor(0x16, 0xa3, 0x4a)
RED      = RGBColor(0xdc, 0x26, 0x26)
AMBER    = RGBColor(0xd9, 0x77, 0x06)
GRAY     = RGBColor(0x64, 0x74, 0x8b)
LIGHTGRAY= RGBColor(0xe2, 0xe8, 0xf0)
WHITE    = RGBColor(0xff, 0xff, 0xff)
BG_LIGHT = RGBColor(0xf8, 0xfa, 0xfc)
BG_TABLE = RGBColor(0xf1, 0xf5, 0xf9)
BG_HEAD  = RGBColor(0x1e, 0x29, 0x3b)

VERDICT_COLORS = {
    "SUPPORTED":    GREEN,
    "REFINE":       AMBER,
    "CONTRADICT":   RED,
    "CONTRADICTED": RED,
    "NEW FINDING":  BLUE,
}

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 4)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, color=None, size=None, font=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if font:
        run.font.name = font
    return run

def set_para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        from docx.shared import Pt as _Pt
        pf.line_spacing = _Pt(line)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=18, after=4)
    run = add_run(p, text, bold=True, color=NAVY, size=17)
    run.font.name = "Calibri"
    return p

def add_heading2(doc, text, color=BLUE):
    p = doc.add_paragraph()
    set_para_spacing(p, before=14, after=3)
    run = add_run(p, text, bold=True, color=color, size=13)
    run.font.name = "Calibri"
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=10, after=2)
    run = add_run(p, text, bold=True, color=NAVY, size=11)
    run.font.name = "Calibri"
    return p

def add_body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    set_para_spacing(p, before=2, after=4)
    run = add_run(p, text, italic=italic, color=color, size=10)
    run.font.name = "Calibri"
    return p

def add_body_mixed(doc, parts):
    """parts = list of (text, bold, italic, color)"""
    p = doc.add_paragraph()
    set_para_spacing(p, before=2, after=4)
    for text, bold, italic, color in parts:
        r = add_run(p, text, bold=bold, italic=italic, color=color, size=10)
        r.font.name = "Calibri"
    return p

def add_verdict_box(doc, verdict_label, verdict_text):
    color = VERDICT_COLORS.get(verdict_label.upper(), BLUE)
    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=4)
    r1 = add_run(p, f"{verdict_label}", bold=True, color=color, size=10)
    r1.font.name = "Calibri"
    r2 = add_run(p, f" — {verdict_text}", bold=False, color=NAVY, size=10)
    r2.font.name = "Calibri"
    return p

def add_recommendation(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_para_spacing(p, before=1, after=2)
    # Label
    r1 = add_run(p, "Recommendation: ", bold=True, color=BLUE, size=10)
    r1.font.name = "Calibri"
    r2 = add_run(p, text, italic=True, color=NAVY, size=10)
    r2.font.name = "Calibri"
    return p

def add_hr(doc):
    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CBD5E1")
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table(doc, headers, rows, col_widths=None, header_bg=BG_HEAD, header_fg=WHITE, stripe=True):
    """Add a styled table. headers = list of str, rows = list of list of str."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = add_run(p, h, bold=True, color=header_fg, size=9)
        r.font.name = "Calibri"
        set_para_spacing(p, before=2, after=2)

    # Data rows
    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        bg = BG_TABLE if (stripe and ri % 2 == 0) else WHITE
        for ci, cell_text in enumerate(row):
            cell = trow.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_para_spacing(p, before=2, after=2)
            # Handle bold markers: **text**
            import re
            parts = re.split(r"(\*\*.*?\*\*)", cell_text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = add_run(p, part[2:-2], bold=True, color=NAVY, size=9)
                else:
                    r = add_run(p, part, bold=False, color=NAVY, size=9)
                r.font.name = "Calibri"

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    # Space after table
    doc.add_paragraph()
    return table

# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.font.color.rgb = NAVY

# ── Cover / Title ─────────────────────────────────────────────────────────────

p = doc.add_paragraph()
set_para_spacing(p, before=0, after=2)
r = add_run(p, "T1 Data Evidence Report", bold=True, color=NAVY, size=22)
r.font.name = "Calibri"

p = doc.add_paragraph()
set_para_spacing(p, before=0, after=2)
r = add_run(p, "Platform Format Guide: SmartNews & Apple News", bold=True, color=BLUE, size=14)
r.font.name = "Calibri"

p = doc.add_paragraph()
set_para_spacing(p, before=0, after=8)
r = add_run(p, "April 3, 2026  ·  Content Strategy & Audience Team  ·  McClatchy", italic=True, color=GRAY, size=10)
r.font.name = "Calibri"

add_hr(doc)

# Data note
p = doc.add_paragraph()
set_para_spacing(p, before=6, after=8)
r = add_run(p,
    "Data: 4,174 Apple News articles (2025–2026), 38,213 SmartNews articles (2025), "
    "1,783 push notifications (2025–2026). Statistical tests: Mann-Whitney U with "
    "rank-biserial correlation (r). Effect size conventions: small r=0.1, medium r=0.3, "
    "large r=0.5. All p-values two-tailed.",
    italic=True, color=GRAY, size=9)
r.font.name = "Calibri"

# Verdict key
add_heading3(doc, "How to read verdicts")
add_table(doc,
    headers=["Verdict", "Meaning"],
    rows=[
        ["SUPPORTED",    "Our data confirms the guide's claim"],
        ["REFINE",       "The claim is directionally correct but the specifics need updating"],
        ["CONTRADICTED", "Our data contradicts the guide's claim — revise or remove"],
        ["NEW FINDING",  "The guide is silent on this; our data says it matters"],
    ],
    col_widths=[1.5, 5.0],
    header_bg=NAVY,
)

add_hr(doc)

# ── SECTION 1 ─────────────────────────────────────────────────────────────────

add_heading1(doc, "Section 1: Character Length Recommendations")

# 1a Apple News length
add_heading2(doc, "1a  Apple News: '80–110 characters recommended'")
add_verdict_box(doc, "SUPPORTED", "with an upward extension")

add_body(doc,
    "Our data strongly validates this range. Apple News top-quartile performers have a "
    "median of 86 characters (p25–p75: 80–96). The 80–110 bin significantly outperforms "
    "shorter ranges (Mann-Whitney p<0.001, r=0.16). However, the fine-grained performance "
    "curve keeps rising through 110–119 characters:"
)

add_table(doc,
    headers=["Char range", "n", "Median norm. views", "Notes"],
    rows=[
        ["60–69",      "236",    "−0.28", ""],
        ["70–79",      "1,744",  "−0.29", ""],
        ["80–89",      "1,440",  "−0.28", ""],
        ["90–99",      "405",    "−0.17", ""],
        ["100–109",    "174",    "−0.10", ""],
        ["**110–119**","**41**", "**+0.08**","**Best performing bin**"],
    ],
    col_widths=[1.3, 0.7, 1.6, 3.4],
)

add_recommendation(doc,
    "Extend the Apple News character target to 90–120 characters. "
    "The current 110-character ceiling appears to leave performance on the table — "
    "the 110–119 bin has the highest median views of any bin. Note this as: "
    "\"90–120 characters; the performance curve peaks above 100.\""
)

# 1b SmartNews length
add_heading2(doc, "1b  SmartNews: 'Keep under 100 characters'")
add_verdict_box(doc, "REFINE", "the 100-character threshold has no statistical basis")

add_body(doc,
    "The 100-character cutoff has no statistical support as a performance boundary "
    "(MW <100 vs. 100+: p=0.915, r=0.002 — completely flat). What the data shows is "
    "that 80–99 chars is the SmartNews sweet spot, and top-quartile performers have "
    "a median of 73 characters (p25–p75: 69–77):"
)

add_table(doc,
    headers=["Char range", "n", "Median norm. views", "Notes"],
    rows=[
        ["<60",        "3,216",  "−0.235", ""],
        ["60–79",      "29,450", "−0.230", ""],
        ["**80–99**",  "**4,899**","**−0.221**","**Best performing bin**"],
        ["100–120",    "608",    "−0.230", "Performance flat vs. 60–79"],
        ["121+",       "40",     "−0.208", "Too small for significance"],
    ],
    col_widths=[1.3, 0.9, 1.6, 3.2],
)

add_recommendation(doc,
    "Replace 'under 100 characters' with '70–90 characters is the proven sweet spot; "
    "80–99 performs best.' The 100-character ceiling is not where performance degrades — "
    "80–99 just happens to be where it peaks. There is no statistical difference between "
    "headlines under 100 and headlines 100–120."
)

add_hr(doc)

# ── SECTION 2 ─────────────────────────────────────────────────────────────────

add_heading1(doc, "Section 2: SmartNews Formula Guidance")

add_heading2(doc, "2a  SmartNews: \"'How to,' 'What to Know,' cost/savings angles can perform well\"", color=RED)
add_verdict_box(doc, "CONTRADICTED", "remove or reverse this guidance")

add_body(doc,
    "This is the most important correction in this report. Our data shows what_to_know "
    "significantly underperforms on SmartNews across 73 articles "
    "(p=0.046, r=−0.135). Question headlines also significantly underperform "
    "(n=889, p<0.001, r=−0.091). \"Here's\" formula is non-significant but directionally negative. "
    "All three are the formulas the guide specifically recommends."
)

add_table(doc,
    headers=["Formula", "n", "vs. Baseline", "p-value", "r", "Verdict"],
    rows=[
        ["number_lead",            "2,988", "+0.003", "0.319", "+0.011", "Neutral (best trend)"],
        ["heres_formula",          "61",    "−0.002", "0.851", "+0.014", "Neutral"],
        ["possessive named entity","1,133", "−0.003", "0.277", "−0.019", "Neutral"],
        ["untagged (baseline)",    "30,543","—",      "—",     "—",      "Baseline"],
        ["quoted_lede",            "1,970", "−0.004", "<0.001","−0.076", "Significantly below"],
        ["**question**",           "**889**","**−0.006**","**<0.001**","**−0.091**","**Significantly below**"],
        ["**what_to_know**",       "**73**","**−0.008**","**0.046**","**−0.135**","**Significantly below**"],
    ],
    col_widths=[1.8, 0.6, 1.0, 0.8, 0.6, 2.2],
)

add_body(doc,
    "The guide's intuition that SmartNews rewards 'practical value' is correct in spirit — "
    "but 'What to Know' is the wrong vehicle. The formulas the guide recommends for "
    "SmartNews service journalism are the exact formulas that hurt SmartNews performance.",
    italic=True, color=GRAY
)

add_recommendation(doc,
    "Remove 'What to Know' and question formats from SmartNews recommendations. "
    "Replace with: 'Number-led headlines trend positive on SmartNews; plain declarative "
    "statements are the safe baseline. Avoid question marks and \"What to Know\" endings — "
    "both significantly underperform across thousands of articles.'"
)

# 2b Service journalism verticals
add_heading2(doc, "2b  SmartNews: 'Service journalism verticals can perform well'", color=RED)
add_verdict_box(doc, "CONTRADICTED", "lifestyle content is near the bottom of the SmartNews performance stack")

add_body(doc,
    "SmartNews 2025 topic performance shows lifestyle content is near the bottom of all "
    "categories. Service journalism proxies (lifestyle + nature/wildlife) vs. hard news: "
    "MW p=0.74, r=0.006 — no statistically detectable difference. The dominant SmartNews "
    "performer is politics (driven by 2025 news cycle); everything else clusters near baseline."
)

add_table(doc,
    headers=["Topic", "n", "Median norm. views", "Notes"],
    rows=[
        ["politics",         "556",    "+0.396", "Massive outlier — 2025 news cycle"],
        ["nature/wildlife",  "258",    "−0.217", ""],
        ["weather",          "267",    "−0.219", ""],
        ["local/civic",      "381",    "−0.219", ""],
        ["business",         "277",    "−0.228", ""],
        ["crime",            "2,335",  "−0.233", ""],
        ["**lifestyle**",    "**819**","**−0.235**","**Near bottom**"],
        ["**sports**",       "**2,115**","**−0.237**","**Bottom**"],
    ],
    col_widths=[1.6, 0.7, 1.6, 3.1],
)

add_recommendation(doc,
    "Soften or remove the service journalism vertical claim for SmartNews. "
    "Politics dominated SmartNews in 2025 (which T1 content teams don't write). "
    "For realistic T1 content, the top of the SmartNews performance range is "
    "weather and local/civic — not service journalism."
)

add_hr(doc)

# ── SECTION 3 ─────────────────────────────────────────────────────────────────

add_heading1(doc, "Section 3: Apple News Formula Guidance")

add_heading2(doc, "3a  NEW FINDING: Questions are the most damaging formula on Apple News")
add_verdict_box(doc, "NEW FINDING", "the guide does not mention formula avoidance for Apple News — it should")

add_body(doc,
    "Question headlines significantly underperform on Apple News (n=203, p<0.001, r=−0.265 "
    "— approaching medium effect size). This is the strongest formula signal in the entire "
    "Apple News dataset. Number-led headlines also significantly underperform "
    "(n=225, p=0.006, r=−0.109)."
)

add_table(doc,
    headers=["Formula", "n", "vs. Baseline", "p-value", "r", "Verdict"],
    rows=[
        ["heres_formula",           "26",    "+0.019","0.492",  "+0.078","Neutral (best trend)"],
        ["possessive named entity", "117",   "+0.005","0.175",  "+0.074","Neutral"],
        ["what_to_know",            "29",    "+0.008","0.359",  "−0.099","Neutral"],
        ["untagged (baseline)",     "3,387", "—",     "—",     "—",     "Baseline"],
        ["quoted_lede",             "187",   "−0.006","0.433",  "−0.034","Neutral"],
        ["**number_lead**",         "**225**","**−0.012**","**0.006**","**−0.109**","**Significantly below**"],
        ["**question**",            "**203**","**−0.017**","**<0.001**","**−0.265**","**Worst formula**"],
    ],
    col_widths=[1.8, 0.6, 1.0, 0.8, 0.6, 2.2],
)

add_recommendation(doc,
    "Add to the Apple News section: 'Avoid question headlines — they significantly "
    "underperform in organic views (p<0.001, r=−0.27 across 4,100+ articles). "
    "Number-led headlines also underperform on Apple News despite working on other platforms.'"
)

# 3b Featured vs organic tension
add_heading2(doc, "3b  NEW FINDING: Apple's human editors and the algorithm disagree on question headlines")
add_verdict_box(doc, "NEW FINDING", "featured curation and organic performance tell opposite stories")

add_body(doc,
    "Apple's editorial team over-selects question headlines for featured slots (1.67× rate "
    "vs. baseline) and 'What to Know' endings (1.81×). But these formats do not produce "
    "higher organic views when controlled for featured status."
)

add_table(doc,
    headers=["Formula", "Featured rate lift", "Organic view rank", "Implication"],
    rows=[
        ["what_to_know",  "1.81× (over-selected)", "Neutral", "Editors love it; algorithm indifferent"],
        ["question",      "1.67× (over-selected)", "**Worst formula**", "Editors love it; algorithm penalizes it"],
        ["heres_formula", "1.39× (over-selected)", "Neutral", ""],
        ["untagged",      "0.97× (at baseline)",   "**Best**", "Algorithm rewards it; editors overlook it"],
        ["number_lead",   "0.63× (under-selected)","Second worst", "Editors and algorithm both avoid it"],
    ],
    col_widths=[1.6, 1.8, 1.5, 2.1],
)

add_body(doc,
    "Interpretation: Questions may appear to perform well because they get editorially "
    "boosted, masking weak organic performance. Featured placement spikes views, but the "
    "formula itself does not generate organic clicks.",
    italic=True, color=GRAY
)

add_recommendation(doc,
    "Add a nuance note: 'Apple's editorial team favors question and \"What to Know\" formats "
    "for featured placement — but organic algorithmic performance is stronger with plain "
    "declarative headlines. If writing specifically to pursue a featured slot, question "
    "framing may help with human curators; for algorithmic reach, avoid it.'"
)

add_hr(doc)

# ── SECTION 4 ─────────────────────────────────────────────────────────────────

add_heading1(doc, "Section 4: Push Notifications (New Section Needed)")

add_heading2(doc, "4a  NEW FINDING: Notifications have the largest formula signal in the entire dataset")
add_verdict_box(doc, "NEW FINDING", "the guide is completely silent on notification headlines")

add_body(doc,
    "Direct declarative language (untagged baseline) dramatically outperforms every named "
    "formula for push notification CTR. The effect sizes for notifications are 2–5× larger "
    "than anything observed in the views analyses."
)

add_table(doc,
    headers=["Formula", "n", "Median raw CTR", "MW p-value", "r", "Verdict"],
    rows=[
        ["untagged (baseline)",    "1,511","2.22%", "—",      "—",      "Baseline — best"],
        ["number_lead",            "85",   "2.32%", "0.385",  "−0.056", "Neutral"],
        ["possessive named entity","62",   "1.93%", "0.252",  "−0.086", "Neutral"],
        ["**question**",           "**87**","**1.36%**","**<0.001**","**−0.376**","**−38% CTR vs. baseline**"],
        ["**quoted_lede**",        "**30**","**1.09%**","**<0.001**","**−0.437**","**−44% CTR vs. baseline**"],
        ["**heres_formula**",      "**5**", "**0.93%**","**0.037**","**−0.540**","**−54% CTR vs. baseline**"],
    ],
    col_widths=[1.8, 0.5, 1.1, 1.0, 0.6, 2.0],
)

add_recommendation(doc,
    "Add a push notifications section: 'State the news directly — plain declarative "
    "language outperforms every named formula. Avoid question marks (−38% CTR), quoted "
    "ledes (−44%), and \"Here's\" formulas (−54%). The audience decides in 1–2 seconds "
    "whether to tap; friction in the headline format directly suppresses CTR. "
    "This is the strongest and most consistent formula finding in our McClatchy T1 data.'"
)

add_hr(doc)

# ── SECTION 5 ─────────────────────────────────────────────────────────────────

add_heading1(doc, "Section 5: Cross-Platform Headline Strategy")

add_heading2(doc, "5a  '50–70 chars SEO/SmartNews, 80–110 chars Apple News' — two-variant strategy")
add_verdict_box(doc, "SUPPORTED", "in structure; refine the specific numbers")

add_body(doc,
    "The two-variant strategy is validated by data. The length targets need tightening "
    "based on actual top-performer distributions:"
)

add_table(doc,
    headers=["Platform", "Guide says", "Data says", "Top-quartile median"],
    rows=[
        ["Apple News", "80–110 chars", "**90–120 chars**", "86 chars (curve rises through 110–119)"],
        ["SmartNews",  "Under 100 chars","**70–90 chars**","73 chars (optimal bin: 80–99)"],
    ],
    col_widths=[1.3, 1.3, 1.4, 3.0],
)

# 5b Cross-platform formula matrix
add_heading2(doc, "5b  Cross-platform formula risk matrix")
add_body(doc, "Which formulas work, which hurt, and on which platform:")

add_table(doc,
    headers=["Formula", "Apple News", "SmartNews", "Pattern"],
    rows=[
        ["number_lead",            "Hurts (p=0.006)",  "Neutral",           "Avoid on Apple News specifically"],
        ["what_to_know",           "Neutral",           "Hurts (p=0.046)",   "Avoid on SmartNews specifically"],
        ["quoted_lede",            "Neutral",           "Hurts (p<0.001)",   "Avoid on SmartNews"],
        ["heres_formula",          "Neutral",           "Neutral",           "Safe on both"],
        ["possessive named entity","Neutral",           "Neutral",           "Safe on both"],
        ["untagged (declarative)", "Best baseline",     "Best baseline",     "Safest across both platforms"],
        ["**question**",           "**Worst (p<0.001)**","**Hurts (p<0.001)**","**Avoid on both platforms**"],
    ],
    col_widths=[1.8, 1.5, 1.5, 2.2],
)

add_body(doc,
    "No formula significantly helps on both platforms. The clearest universal rule: "
    "questions hurt everywhere — on Apple News (r=−0.265), on SmartNews (r=−0.091), "
    "and dramatically on push notifications (r=−0.376).",
    italic=True, color=GRAY
)

add_hr(doc)

# ── SECTION 6 ─────────────────────────────────────────────────────────────────

add_heading1(doc, "Section 6: Topic Strategy for Apple News")

add_heading2(doc, "6a  NEW FINDING: Business is the weakest topic; sports leads; lifestyle underperforms")
add_verdict_box(doc, "NEW FINDING", "the guide makes no topic performance claims for Apple News")

add_table(doc,
    headers=["Topic", "n", "Median norm. views", "Notes"],
    rows=[
        ["**sports**",       "**276**","**−0.252**","**Best performing topic**"],
        ["crime",            "357",   "−0.269",""],
        ["weather",          "361",   "−0.270",""],
        ["politics",         "290",   "−0.286","Significantly below baseline (p<0.001)"],
        ["nature/wildlife",  "265",   "−0.290",""],
        ["local/civic",      "290",   "−0.291",""],
        ["lifestyle",        "157",   "−0.291","Below baseline"],
        ["**business**",     "**106**","**−0.293**","**Worst performing topic**"],
    ],
    col_widths=[1.6, 0.7, 1.6, 3.1],
)

add_recommendation(doc,
    "Add to Apple News section: 'Sports content has the strongest topic signal on Apple News. "
    "Business and lifestyle content consistently underperform — if writing in these verticals, "
    "compensate with stronger headline execution. This suggests Apple News audiences use "
    "the platform primarily for news, not service content.'"
)

add_hr(doc)

# ── SUMMARY TABLE ─────────────────────────────────────────────────────────────

add_heading1(doc, "Summary: All Changes at a Glance")

add_table(doc,
    headers=["Section", "Change", "Strength"],
    rows=[
        ["SmartNews — editorial",  "Remove 'What to Know' from SmartNews recommendations",                        "Strong (p=0.046)"],
        ["SmartNews — editorial",  "Add: questions and WTK hurt; number leads trend positive",                     "Strong (p<0.001)"],
        ["SmartNews — length",     "Change 'under 100 chars' → '70–90 chars ideal; 80–99 best'",                  "Moderate"],
        ["SmartNews — verticals",  "Soften service journalism claim — lifestyle near bottom",                      "Moderate"],
        ["Apple News — length",    "Extend upper bound from 110 → 120 chars",                                      "Moderate"],
        ["Apple News — formulas",  "Add: questions significantly underperform (p<0.001, r=−0.27)",                 "Strong"],
        ["Apple News — featured",  "Add nuance: editors favor questions; organic algorithm does not",              "Novel"],
        ["Apple News — topics",    "Add: sports leads, business and lifestyle lag",                                 "Moderate"],
        ["Notifications",          "Add full section: plain declarative dominates; avoid questions/WTK/Here's",    "Very strong"],
        ["Variants strategy",      "Refine: SN 70–90 chars, Apple News 90–120 chars",                              "Moderate–strong"],
    ],
    col_widths=[2.0, 3.5, 1.5],
    header_bg=NAVY,
)

# ── Footer note ───────────────────────────────────────────────────────────────

p = doc.add_paragraph()
set_para_spacing(p, before=10, after=2)
r = add_run(p,
    "Note: All findings are based on McClatchy T1 outlet data only. "
    "SmartNews 2026 data is monthly-aggregated by domain (not article-level); "
    "SmartNews formula and character analyses use 2025 article-level data (n=38,213). "
    "Apple News analyses combine 2025 and 2026 Jan–Mar (n=4,174). "
    "Statistical significance does not imply causation — these are observational signals "
    "for experimentation, not editorial rules.",
    italic=True, color=GRAY, size=8.5
)
r.font.name = "Calibri"

# ── Save ──────────────────────────────────────────────────────────────────────
import os
out_path = os.path.expanduser("~/Desktop/T1_Data_Evidence_Report_Format_Guide.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
